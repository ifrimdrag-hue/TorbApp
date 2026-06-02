"""Word handout — brief executiv pentru director."""

from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import theme as T


def _hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, color_hex: str):
    """Seteaza fundal pe o celula de tabel."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tc_pr.append(shd)


def _h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = T.FONT_TITLE
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(T.NAVY)
    return p


def _h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = T.FONT_TITLE
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(T.NAVY)
    return p


def _h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = T.FONT_BODY
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(T.DARK_GRAY)
    return p


def _body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = T.FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = _hex_to_rgb(color or T.DARK_GRAY)
    if italic:
        run.italic = True
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = T.FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = _hex_to_rgb(T.DARK_GRAY)
    if not p.runs[0].text:
        p.add_run(text)


def _styled_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(h)
        run.font.name = T.FONT_BODY
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = _hex_to_rgb(T.WHITE)
        _set_cell_bg(hdr_cells[i], T.NAVY)

    # Data rows
    for r_idx, row_data in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ""
            para = cells[c_idx].paragraphs[0]
            run = para.add_run(str(val))
            run.font.name = T.FONT_BODY
            run.font.size = Pt(10)
            run.font.color.rgb = _hex_to_rgb(T.DARK_GRAY)
        if r_idx % 2 == 0:
            for c in cells:
                _set_cell_bg(c, T.LIGHT_GRAY)
    return table


def build_docx(campaigns: list[dict]) -> bytes:
    doc = Document()

    # Margini
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ─── Header / Titlu ───
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Plan Campanii Mai 2026")
    run.font.name = T.FONT_TITLE
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(T.NAVY)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Marketing & E-commerce — Brief executiv pentru sedinta")
    sub_run.font.name = T.FONT_BODY
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = _hex_to_rgb(T.MID_GRAY)
    sub_run.italic = True

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Data: {date.today().strftime('%d %B %Y')}")
    meta_run.font.name = T.FONT_BODY
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = _hex_to_rgb(T.MID_GRAY)

    doc.add_paragraph()  # spacer

    # ─── 1. Sumar executiv ───
    _h1(doc, "1. Sumar executiv")
    _body(doc,
          "Propun 3 campanii pentru luna mai 2026, complementare si fara "
          "canibalizare. Bugetul total ads: 1.000 RON (200 EUR). Obiectiv: "
          "mentinere/crestere vanzari online, focus pe marja contributie.")
    doc.add_paragraph()

    _h3(doc, "Cele 3 campanii")
    _bullet(doc, "Mama Mea — Pachet Cadou (gifting Ziua Mamei, 30 apr → 3 mai, 400 RON)")
    _bullet(doc, "Sezonul Schimba — Refresh Primavara (lifestyle, 5 → 25 mai, 300 RON)")
    _bullet(doc, "Verile Usor — Pre-Summer Healthy (sezonier, 20 mai → 10 iun, 300 RON)")
    doc.add_paragraph()

    # ─── 2. Context comercial ───
    _h1(doc, "2. Context — De ce mai 2026?")
    _body(doc,
          "Mai este una dintre cele mai dinamice luni comerciale din anul calendaristic "
          "in Romania. Trei momente cu potential demonstrat istoric:")
    doc.add_paragraph()

    _styled_table(doc,
        ["Moment", "Data", "Tip oportunitate"],
        [
            ["Ziua Mamei", "duminica 3 mai", "Gifting (peak comercial, comparabil cu Martisor)"],
            ["Tranzitie primavara→vara", "5-25 mai", "Lifestyle, refresh, terase deschise"],
            ["Pre-vara 'light'", "20 mai → iunie", "Sanatate, fara zahar, hidratare"],
        ],
    )
    doc.add_paragraph()

    # ─── 3. Detalii pe campanie ───
    _h1(doc, "3. Detalii pe campanie")
    doc.add_paragraph()

    type_labels = {"promo": "Promo", "gifting": "Gifting", "lansare": "Lansare", "sezonier": "Sezonier", "giveaway": "Giveaway"}
    channel_labels = {"shopify": "Shopify", "emag": "eMAG", "instagram": "Instagram", "facebook": "Facebook"}

    for idx, c in enumerate(campaigns, start=1):
        _h2(doc, f"{idx}. {c.get('name', '')}")

        info_rows = [
            ["Tip", type_labels.get(c.get("type"), c.get("type", ""))],
            ["Perioada", f"{c.get('date_start','')} → {c.get('date_end','')}"],
            ["Canale", ", ".join(channel_labels.get(ch, ch) for ch in c.get("channels", []))],
            ["Mecanica", c.get("mechanic", "")],
            ["Buget alocat", f"{int(c.get('budget_alloc') or 0)} RON"],
            ["Produse incluse", str(len(c.get("products", []) or []))],
        ]
        _styled_table(doc, ["Camp", "Valoare"], info_rows)
        doc.add_paragraph()

        if c.get("notes"):
            _h3(doc, "Note strategice")
            _body(doc, c["notes"])
            doc.add_paragraph()

        # Top produse
        prods = (c.get("products") or [])[:6]
        if prods:
            _h3(doc, "Produse principale")
            for p in prods:
                _bullet(doc, f"{p.get('name', p.get('sku', ''))}" +
                        (f" — necesar {p.get('qty_needed')} buc" if p.get("qty_needed") else ""))
            doc.add_paragraph()

    # ─── 4. Buget consolidat ───
    _h1(doc, "4. Buget consolidat")
    doc.add_paragraph()

    total_alloc = sum(c.get("budget_alloc") or 0 for c in campaigns)
    total_spent = sum(c.get("budget_spent") or 0 for c in campaigns)

    rows_buget = [[c.get("name", ""), f"{int(c.get('budget_alloc') or 0)} RON",
                   f"{int(c.get('budget_spent') or 0)} RON"] for c in campaigns]
    rows_buget.append(["TOTAL", f"{int(total_alloc)} RON", f"{int(total_spent)} RON"])

    _styled_table(doc, ["Campanie", "Alocat (ads)", "Cheltuit"], rows_buget)
    doc.add_paragraph()

    _body(doc, "Buget producere creativa (foto, video, design) — separat: estimat 1.500-2.500 RON. "
               "Asseturi reutilizabile peste 3 campanii (amortizare buna).", italic=True)
    doc.add_paragraph()

    # ─── 5. KPI tinta ───
    _h1(doc, "5. KPI tinta agregate")
    _body(doc, "Rezultate asteptate la sfarsitul lunii (presupunand executia conform planului):")
    doc.add_paragraph()

    _styled_table(doc, ["Indicator", "Tinta"], [
        ["Comenzi B2C totale", "350-450"],
        ["Revenue B2C atribuibil", "25.000 - 40.000 RON"],
        ["AOV mediu", "+15-25% vs aprilie (datorita bundle-urilor)"],
        ["Contribution margin B2C", "mentinut peste 28%"],
        ["Email subscribers nou", "+200-400"],
        ["Lichidare stoc Torras (campania C)", "60-70% din stocul curent"],
    ])
    doc.add_paragraph()

    # ─── 6. Riscuri & mitigari ───
    _h1(doc, "6. Riscuri & mitigari")
    doc.add_paragraph()
    _styled_table(doc, ["Risc", "Probabilitate", "Mitigare"], [
        ["Stoc insuficient Tea Book Mama Mea", "Medie", "Verific stoc prin Hub inainte de lansare; limitez bundle"],
        ["Canibalizare intre B si C", "Mica", "Audiente diferite (lifestyle vs health)"],
        ["Vreme rece prelungita", "Mica-Medie", "Pivotare mesaj la 'indoor coziness'"],
        ["Foto-sesiunea intarziata", "Medie", "Programare imediat dupa aprobare; buffer 2 zile"],
        ["App 2+1 Shopify nu functioneaza", "Identificat", "Plan B: cod discount manual"],
    ])
    doc.add_paragraph()

    # ─── 7. Decizii de luat ───
    _h1(doc, "7. Decizii necesare in sedinta")
    _body(doc, "Pentru lansare conform calendarului, am nevoie de aprobare pe urmatoarele puncte:")
    doc.add_paragraph()

    _bullet(doc, "Aprobare buget total ads: 1.000 RON")
    _bullet(doc, "Aprobare buget producere creativa: 1.500-2.500 RON (sau utilizare librarie existenta)")
    _bullet(doc, "Confirmare stoc suficient Basilur Tea Book Vol I/II/III pentru bundle-ul de Ziua Mamei")
    _bullet(doc, "Coordonare cu KAM pentru campania C (potentiali clienti B2B: gym-uri, cafenele healthy)")
    _bullet(doc, "Validare ROI istoric Ziua Mamei 2025 pentru calibrare asteptari")
    doc.add_paragraph()

    # ─── 8. Next steps ───
    _h1(doc, "8. Next steps imediate (post-aprobare)")
    doc.add_paragraph()
    _bullet(doc, "Lansare Mama Mea — 30 aprilie")
    _bullet(doc, "Foto-sesiune programata — 1-2 mai")
    _bullet(doc, "Setup ads pe Meta — pana la 1 mai 18:00")
    _bullet(doc, "Listare bundle pe Shopify — 30 aprilie")
    _bullet(doc, "Check-in saptamanal cu directorul comercial — vineri")

    # ─── Footer mic ───
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Document generat din Hub Automatizari — Marketing E-commerce")
    fr.font.size = Pt(8)
    fr.font.color.rgb = _hex_to_rgb(T.MID_GRAY)
    fr.italic = True

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
